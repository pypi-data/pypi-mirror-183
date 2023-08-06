from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START,WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Cm, Pt, RGBColor



contents=[
    {
        "block":"section",
        "style":{}
    },
    {
        "block":"styles",
        "style":{}
    },
    {
        "block":"heading",
        "level":1,
        "text":"Submission Letter",
        "style":{}
    },
    {   
        "block":"paragraph",
        "text":"Underline is a bit of a special case. It is a hybrid of a tri-state property and an enumerated value property. True means single underline, by far the most common. False means no underline, but more often {} is the right choice if no underlining is wanted since it is rare to",
        "style":{}
    },
    {
        "block":"heading",
        "level":2,
        "text":"ParagraphFormat",
        "style":{}
    },
    {
        "block":"paragraph",
        "text":"Here’s an example of how you would create a paragraph style having hanging indentation of 1/4 inch, 12 points spacing above, and widow/orphan control:",
        "style":{
            "first_line_indent":{
                "unit":"inch",
                "value":0.25
            }
        }
    },
    {
        "block":"heading",
        "level":3,
        "text":"ParagraphFormat",
        "style":{}
    },
    {
        "block":"paragraph",
        "text":"Here’s an example of how you would create a paragraph style having hanging indentation of 1/4 inch, 12 points spacing above, and widow/orphan control:",
        "style":{
            "left_indent":{
                "unit":"inch",
                "value":0.5
            }
        }
    },
    {
        "block":"run",
        "text":"Jacky",  
        "style":{"name":"Arial"}
    },
    {
        "block":"table",
        "text":[
            ['name','sex','age','education'],
            ['jacky','Male','52','Master'],
            ['dan','male','11','elementory']
        ],  
        "style":"Medium Shading 1 Accent 2"
    }
]   

class Word():
    default_styles={
        "Normal":{
            "name":"Calibri",
            "size":Pt(11),
            "color":RGBColor(0,0,0)
        },
        "Heading 1":{
            "name":"Arial Black",
            "size":Pt(36),
            "color":RGBColor(0,0,125)
        },
        "Heading 2":{
            "name":"Arial Black",
            "size":Pt(24),
            "color":RGBColor(0,125,0)
        },
        "Heading 3":{
            "name":"Arial Black",
            "size":Pt(18),
            "color":RGBColor(0,125,0)
        }
        
    }
    default_paragraph_format={
        "name":"Indent",
        "left_indent":{
            "unit":"inch",
            "value":0
            },
        "first_line_indent":{
            "unit":"inch",
            "value":0
        },
        "space_before":{
            "unit":"pt",
            "value":12
        },
        "space_after":{
            "unit":"pt",
            "value":18
        },
        "widow_control":True
    }
    default_run_format={
        "name":'Calibri',
        "size":Pt(36),
        "italic":True,
        "bold":True,
        "underline":True,
        "color":RGBColor(0,0,255)
    }

    def __init__(self,contents,fileName):
        self.content=contents
        self.fileName=fileName
        self.document=Document()

    def __getUnitValue(self,format,key):
        if  format[key].get("unit","pt")=='pt':
            return Pt(format[key].get("value",None))
        elif format[key].get("unit","pt")=='inch':
            return Inches(format[key].get("value",None))
        elif format[key].get("unit","pt")=='cm':
            return Cm(format[key].get("value",None))
        else:
            return None
    # 2D list
    def __creatTable(self,dataList,style):
        rows=len(dataList)
        cols=len(dataList[0])
        table=self.document.add_table(rows=rows,cols=cols)
        table.style=style
        for row in range(0,rows):
            for col in range(0,cols):
                cell=table.cell(row,col)
                cell.text=dataList[row][col]

    def create(self):
        lastParagraph=None
        for content in contents:
            if content['block']=='styles':
                # if content has its own format, overwrite the counterpart default format
                styleData={**Word.default_styles,**content.get('styles',{})}
                print("document style: ", styleData)
                for k, v in Word.default_styles.items():
                    font=self.document.styles[k].font
                    font.name=v.get('name',None)
                    font.size=v.get('size',None)
                    font.color.rgb=v.get('color',None)
            elif content['block']=='paragraph':
                format={**Word.default_paragraph_format,**content.get('style',{})}
                print('paragraphp format: ',format)
                text=content.get('text',"")
                paragrah=self.document.add_paragraph(text)
                # paragraph formatting
                paragraph_format=paragrah.paragraph_format
                paragraph_format.space_before = self.__getUnitValue(format,"space_before")
                paragraph_format.space_after=self.__getUnitValue(format,"space_after")
                paragraph_format.left_indent=self.__getUnitValue(format,"left_indent")
                paragraph_format.first_line_indent=self.__getUnitValue(format,"first_line_indent")
                paragraph_format.widow_control =format.get("widow_control",True)
                lastParagraph=paragrah

            elif content['block']=='heading':
                self.document.add_heading(content.get('text',''),level=content.get("level",None))
            elif content['block']=='table':
                self.__creatTable(content.get("text",[]),content.get("style","Normal"))
            elif content['block']=='run':
                run=lastParagraph.add_run(content.get("text",''))
                format={**Word.default_run_format,**content.get('style',{})}
                print("run format: ",format)
                font=run.font
                font.name=format.get("name","Calibri")
                font.size=format.get("size",None)
                font.italic=format.get("italic",None)
                font.bold=format.get("bold",None)
                font.underline=format.get("underline",None)
                font.color.rgb=format.get("color",RGBColor(0,0,255))
            else:
                pass 
        self.document.save(self.fileName)

w=Word(contents,"jacky.docx")
w.create()
